from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
MD = BASE / "20260625_001_paper_final.md"
OUT = BASE / "20260625_001_大模型在智慧交通调度中的应用研究.docx"
OUT_FIXED = BASE / "20260625_001_大模型在智慧交通调度中的应用研究_表格修正版.docx"
OUT_IMAGE_FIXED = BASE / "20260625_001_大模型在智慧交通调度中的应用研究_图片修正版.docx"

def set_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)

def style_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, first=0, line=None, before=0, after=0):
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Pt(first) if first else None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line

def add_para(doc, text, font="宋体", size=12, bold=False, first=24, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    style_para(p, align=align, first=first, line=1.0)
    r = p.add_run(text)
    set_font(r, font, size, bold)
    return p

def clean_inline(text):
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()

def add_table(doc, lines):
    rows = []
    for line in lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue
        rows.append([c.strip() for c in line.strip().strip("|").split("|")])
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = [Cm(3.0), Cm(3.4), Cm(2.7), Cm(3.0), Cm(4.0)]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            if j < len(widths):
                cell.width = widths[j]
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, "宋体", 10, bold=(i == 0))
    doc.add_paragraph()

def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.175)
        section.right_margin = Cm(3.175)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    in_refs = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            add_para(doc, line[2:].strip(), "黑体", 16, True, first=0)
            i += 1
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            in_refs = "参考文献" in title
            if title.startswith("表"):
                add_para(doc, title, "宋体", 10.5, True, first=0, align=WD_ALIGN_PARAGRAPH.CENTER)
                i += 1
                continue
            add_para(doc, title, "黑体", 14, False, first=0)
            i += 1
            continue
        if line.startswith("### "):
            add_para(doc, line[4:].strip(), "黑体", 14, False, first=0)
            i += 1
            continue
        img = re.match(r"!\[([^\]]+)\]\(([^)]+)\)", line)
        if img:
            alt, rel = img.group(1), img.group(2)
            path = BASE / rel
            if alt.startswith("表"):
                add_para(doc, alt, "宋体", 10.5, True, first=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            if path.exists():
                r.add_picture(str(path), width=Inches(5.8))
            if not alt.startswith("表"):
                add_para(doc, alt, "宋体", 10.5, False, first=0, align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines)
            continue
        text = clean_inline(line)
        if text:
            if text.startswith("关键词："):
                p = doc.add_paragraph()
                style_para(p, first=24, line=1.0)
                r1 = p.add_run("关键词：")
                set_font(r1, "黑体", 12, False)
                r2 = p.add_run(text.replace("关键词：", "", 1))
                set_font(r2, "宋体", 12, False)
            elif text.startswith("摘要："):
                p = doc.add_paragraph()
                style_para(p, first=24, line=1.0)
                r1 = p.add_run("摘要：")
                set_font(r1, "黑体", 12, False)
                r2 = p.add_run(text.replace("摘要：", "", 1))
                set_font(r2, "宋体", 12, False)
            elif in_refs:
                add_para(doc, text, "宋体", 12, False, first=0)
            else:
                add_para(doc, text, "宋体", 12, False, first=24)
        i += 1
    output = OUT_IMAGE_FIXED
    doc.save(output)
    print(output)

if __name__ == "__main__":
    main()
